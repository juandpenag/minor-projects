import emoji
import csv
import os
import shutil
from docx import Document # type: ignore
from docx.enum.text import WD_ALIGN_PARAGRAPH
import pandas as pd
import sys
import time
import re

# Global variables
library_path = "E:\Library" # Main folder
library_csv = "E:\Library\Library.csv" # CSV file
try: #DataFrame to work with
    df = pd.read_csv(library_csv, encoding='ISO-8859-1')  
except FileNotFoundError as e:
    print(f"The CSV file was not found at {library_csv}. Error: {e}")  
except pd.errors.EmptyDataError  as e:
    print("The CSV file is empty: {e}")
except pd.errors.ParserError as e:
    print(f"Error parsing the CSV: {e}")
except Exception as e:
    print(f"Unexpected error: {e}")

def main(): # Welcome / Add or Search?
    print(emoji.emojize('This is Academy! 📚 📖'))

    while True:
        do = input("What do you want to do? Add or Search? (A/S) ").strip().upper()
        if do == "A":
            add()
            break
        elif do == "S":
            search()
            break
        else:
            print("Please, enter a valid choice.")

def add(): # Add "Book"
    
    new_row = { # dict var into which the input is stored
    "Title": input("Title: "), 
    "Authors": input("Authors: "), 
    "Category": input("Category: "), 
    "Subcategory": input("Subcategory: "), 
    "Language": input("Language: "), 
    "Type": input("Type: "), 
    "Resume": "", 
    "File": input("File's path: ").replace('"', '')
    }

    try: # Prepares folder
        file_path = os.path.join(library_path, new_row['Authors']) # Path for the folder in which the file will be stored
        if not os.path.exists(file_path): # Checks if the folder doesn't already exists
            os.makedirs(file_path) # If so, creates the folder
    except PermissionError as e:
        print(f"Permission denied: Unable to create directory at {file_path}")
        sys.exit(1)
    except Exception as e:
        print(f"Unexpected error while preparing the folder: {e}")
        sys.exit(2)
    
    try: # Cuts and pastes file into assigned folder
        moved_file_path = os.path.join(file_path, os.path.basename(new_row['File'])) # Future path of file
        shutil.move(new_row["File"], moved_file_path) # Cuts and pastes the file into its new directory
    except shutil.Error as e:
        print(f"Error moving file: {e}")
        sys.exit(3) 
    except Exception as e:
        print(f"Unexpected error while moving the file: {e}")
        sys.exit(4) 

    try: # Rename the file
        new_file_name = f"{new_row['Title']}.pdf" # Creates new name for file
        new_path = os.path.join(file_path, new_file_name) # Writes down the updated path of file
        os.rename(moved_file_path, new_path) # Changes name and updates path 
    except OSError as e:
        print(f"Error renaming file: {e}")
        sys.exit(5)
    except Exception as e:
        print(f"Unexpected error in the ETL process: {e}")
        sys.exit(6)

    try: # Create "Resume"
        doc = Document() # Creates word doc
        filename = f"{new_row['Title']}'s resume.docx" # Writes down doc's name
        resume_path = os.path.join(file_path, filename) # Writes down resume's path
        new_row["Resume"] = resume_path # Updates resume in new_row

        paragraph = doc.add_paragraph() # Creates the object that allows to add text
        run = paragraph.add_run(f"{new_row['Title']}") # Adds title
        run.bold = True # Makes title bold
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER # Allings title to center
        doc.add_paragraph() # Creates new line in doc
        doc.save(resume_path) # Saves doc in the already formulated path

    except PermissionError:
        print(f"Error: Permission denied when attempting to save the document at {resume_path}.") 
        sys.exit(7)
    except OSError as e: 
        print(f"Unexpected system error when saving the document: {e}")
        sys.exit(8)
    except Exception as e:
        print(f"Unexpected error creating the resume: {e}")
        sys.exit(9)

    try: # Adds new_row into csv through df
        df.loc[len(df)] = new_row # Adds new_row into df
        df.to_csv(library_csv, index=False, encoding='ISO-8859-1') # Updates the csv
    except PermissionError:
        print("Error: Permission denied when trying to write to the CSV file.")
        sys.exit(10)
    except FileNotFoundError:
        print(f"Error: The directory for the CSV file {library_csv} does not exist.")
        sys.exit(11)
    except IOError as e:
        print(f"Error: I/O error while writing to the CSV file: {e}")
        sys.exit(12)
    except UnicodeEncodeError:
        print("Error: Some characters could not be encoded using ISO-8859-1. Consider using UTF-8 or another suitable encoding.")
        sys.exit(13)
    except Exception as e:
        print(f"Unexpected error opening the csv file: {e}")
        sys.exit(14)

    print(emoji.emojize('Run the script again to execute another action.\nThis was Academy! 📚 📖')) # Exit message
    sys.exit(0)

# Search "File"
def search(): 
    # Print Library

    print('Options: ') # Print options
    for column in df.columns:
        if not column in ['File', 'Resume']:
            print(column)
    
    while True: # Check user's choice
        criteria = input("Write your choice: ") 
        if criteria in df.columns: 
            break
        else:
            print("Please, enter a valid option")
            continue

    matches = []  # list to store matches
    while True:
        value = input("What you're looking for? ").strip().lower() # Value to be looked for in the selected option

        for index, row in df.iterrows(): 
            cell_value = str(row[criteria]).strip().lower() # Turns the cell value into a str 
            if re.search(rf".*{re.escape(value)}.*", cell_value, re.IGNORECASE): 
                matches.append(row.to_dict()) # Append the match into matches if a match is found comparing cell_value with value

            # if re.search(r"(.)?+value+(.)?", cell_value, re.IGNORECASE):
                # 

        if matches: # If matches were found
            print("Matches found: ")
            for index, match in enumerate(matches):
                print(f"{index} {match['Title']}. {match['Authors']}. {match['Category']}. {match['Subcategory']}.") # Print matches
            break
        else: # If not
            print("No matches found. Try again.") 
            continue

    while True:
        do = input("Do you want to open a file? (Y/N) ").strip().upper() # Asks user to open a file
        match do:
            case "Y":
                open_file = input("Please, write the title: ") # Asks for the title of the file
                break
            case "N":
                print(emoji.emojize('Run the script again to execute another action.\nThis was Academy! 📚 📖')) # Exit message
                sys.exit(0)
            case _:
                print("Please, enter a valid option.") # Asks for a valid response
                continue

    for row in matches:
        if row['Title'] == open_file: # Searchs file in matches

            try: # Open resume
                print("Opening the resume...")
                os.startfile(row['Resume'])
                time.sleep(5) # Gives time to execute order
            except FileNotFoundError as e: 
                print(f"Resume was not found: {e}")
                sys.exit(15)
            except KeyError as e: 
                print(f"KeyError opening the Resume: {e}")
                sys.exit(16)
            except Exception as e:
                print(f"Unexpected error while opening Resume: {e}")
                sys.exit(17)

            try: # Opens file
                print("Opening the file...")
                os.startfile(row['File']) 
                time.sleep(5) # Gives time to execute order
            except FileNotFoundError as e:
                print(f"The file wasn't found: {e}")
                sys.exit(18)
            except KeyError as e: 
                print(f"KeyError opening the file: {e}")
                sys.exit(19)
            except Exception as e:
                print(f"Unexpected error: {e}")
                sys.exit(20)

    print(emoji.emojize('Run the script again to execute another action.\nThis was Academy! 📚 📖')) # Exit message
    sys.exit(0)

if __name__ == "__main__":
    main()